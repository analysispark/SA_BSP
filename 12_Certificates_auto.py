import docx
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert
from datetime import datetime
import os
import subprocess
import sys

game = '농구'  # 운동종목
time = datetime.today().strftime("%Y.%m.%d")  # 발급날짜
name = 'Park'
birthday = '1987-07-26'
school = '용인대학교'
position = '센터'
sa_KPI_T = '98.82' #Total
sa_KPI_A = '100' #Attack
sa_KPI_D = '97.89' #Defence
sa_KPI_P = '99.22' #Percent

# # for i in range(len(name)):

doc = docx.Document(r'/mnt/d/Work/Python/django/SA_BSP/인증서 양식2.docx')

style = doc.Style['Normal']
style.font.name = '나눔고딕'
style._element.rPr.rFonts.set(qn('w:eastAisa'), '나눔고딕')
style.font.size = docx.shared.Pt(10) # font size

para = doc.add_paragraph()
run = para.add_run('\n\n')
run = para.add_run('제 2022-0001 호\n')
run.font.name = '나눔고딕'
style._element.rPr.rFonts.set(qn('w:eastAisa'), '나눔고딕')
style.font.size = docx.shared.Pt(12) # font size

para = doc.add_paragraph()
run = para.add_run('\n\n')
run = para.add_run('경기력평가지표 증명서')
run.font.name = '나눔고딕'
run.bold = True
style._element.rPr.rFonts.set(qn('w:eastAisa'), '나눔고딕')
style.font.size = docx.shared.Pt(20) # font size
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

para = doc.add_paragraph()
run = para.add_run('\n\n')
run = para.add_run('종    목:' + game + '\n')
run.font.name = '나눔고딕'
style._element.rPr.rFonts.set(qn('w:eastAisa'), '나눔고딕')
style.font.size = docx.shared.Pt(12) # font size

run = para.add_run('소    속:' + school + '\n')
run.font.name = '나눔고딕'
style._element.rPr.rFonts.set(qn('w:eastAisa'), '나눔고딕')
style.font.size = docx.shared.Pt(12) # font size

para = doc.add_paragraph()
run = para.add_run('\n\n')
run = para.add_run('성    명:' + name + '\n')
run.font.name = '나눔고딕'
style._element.rPr.rFonts.set(qn('w:eastAisa'), '나눔고딕')
style.font.size = docx.shared.Pt(12) # font size

run = para.add_run('생년월일:' + birthday + '\n')
run.font.name = '나눔고딕'
style._element.rPr.rFonts.set(qn('w:eastAisa'), '나눔고딕')
style.font.size = docx.shared.Pt(12) # font size

run = para.add_run('포지션  :' + position + '\n')
run.font.name = '나눔고딕'
style._element.rPr.rFonts.set(qn('w:eastAisa'), '나눔고딕')
style.font.size = docx.shared.Pt(12) # font size


para = doc.add_paragraph()
run = para.add_run('\n\n')
run = para.add_run('경기력평가지표:' + sa_KPI_T + '\n')
run.font.name = '나눔고딕'
style._element.rPr.rFonts.set(qn('w:eastAisa'), '나눔고딕')
style.font.size = docx.shared.Pt(20) # font size
run = para.add_run('공격지표:' + sa_KPI_A + '\n')
run.font.name = '나눔고딕'
style._element.rPr.rFonts.set(qn('w:eastAisa'), '나눔고딕')
style.font.size = docx.shared.Pt(20) # font size
run = para.add_run('수비지표:' + sa_KPI_D + '\n')
run.font.name = '나눔고딕'
style._element.rPr.rFonts.set(qn('w:eastAisa'), '나눔고딕')
style.font.size = docx.shared.Pt(20) # font size
run = para.add_run('성공률지표:' + sa_KPI_P + '\n')
run.font.name = '나눔고딕'
style._element.rPr.rFonts.set(qn('w:eastAisa'), '나눔고딕')
style.font.size = docx.shared.Pt(20) # font size

para = doc.add_paragraph()
run = para.add_run('\n\n\n')
run = para.add_run(time)
run.font.name = '나눔고딕'
style._element.rPr.rFonts.set(qn('w:eastAisa'), '나눔고딕')
style.font.size = docx.shared.Pt(20) # font size
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

para = doc.add_paragraph()
run = para.add_run('\n\n\n')
run = para.add_run('용인대학교 산학협력단')
run.font.name = '나눔고딕'
run.bold = True
style._element.rPr.rFonts.set(qn('w:eastAisa'), '나눔고딕')
style.font.size = docx.shared.Pt(20) # font size
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save('평가지표 증명서 자동생성 ' + name + '.docx')

#convert('평가지표 증명서 자동생성 ' + name + '.docx',
#        '평가지표 증명서 자동생성 ' + name + '.pdf')
import re

def convert_to(folder, source, timeout=None):
    args = [libreoffice_exec(), '--headless', '--convert-to', 'pdf', '--outdir', folder, source]

    process = subprocess.run(args, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=timeout)
    filename = re.search('-> (.*?) using filter', process.stdout.decode())

    return filename.group(1)

def libreoffice_exec():
    # TODO: Provide support for more platforms
    if sys.platform == 'darwin':
        return '/usr/bin/soffice'
    return '/usr/bin/soffice'

convert_to('/mnt/d/Work/Python/django/SA_BSP/', "평가지표 증명서 자동생성 " + name + '.docx', timeout=15)

#install
#pip install python-docx
#sudo add-apt-repository ppa:libreoffice/pp
#sudo apt-get update
#sudo apt-get upgrade
#sudo apt install libreoffice
#sudo apt install libreoffice-java-common