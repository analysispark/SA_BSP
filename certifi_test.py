import docx 
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx2pdf import convert
from datetime import datetime
import os
import subprocess
import sys


doc = docx.Document(r'/mnt/d/Work/Python/django/SA_BSP/인증서 양식2.docx')
#doc = docx.Document()

game = '농구'  # 운동종목
time = datetime.today().strftime("%Y.%m.%d")  # 발급날짜
name = '신승윤'
birthday = '1957-03-02'
school = '용인대학교'
position = '가드'

sa_KPI_A = '46.16' #Attack
sa_KPI_D = '46.70' #Defence
sa_KPI_P = '45.43' #Percent

para = doc.add_paragraph()
run = para.add_run('\n\n')
run = para.add_run('제 2023-0001 호\n')
run.font.name = '나눔고딕'
run.font.size = Pt(12)
para.alignment = WD_ALIGN_PARAGRAPH.LEFT

para = doc.add_paragraph()
run = para.add_run('\n\n')
run = para.add_run('경기력평가지표 증명서')
run.font.name = '나눔고딕'
run.bold = True
run.font.size = Pt(20)
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

para = doc.add_paragraph()
run = para.add_run('\n\n')
run = para.add_run('종    목:' + game + '\n')
run.font.name = '나눔고딕'
run.font.size = Pt(15) # font size

run = para.add_run('소    속:' + school + '\n')
run.font.name = '나눔고딕'
run.font.size = Pt(15) # font size

para = doc.add_paragraph()
run = para.add_run('\n')
run = para.add_run('성    명:  ' + name + '\n')
run.font.name = '나눔고딕'
run.font.size = Pt(15) # font size

run = para.add_run('생년월일:  ' + birthday + '\n')
run.font.name = '나눔고딕'
run.font.size = Pt(15) # font size

run = para.add_run('포지션  :  ' + position + '\n')
run.font.name = '나눔고딕'
run.font.size = Pt(15) # font size


para = doc.add_paragraph()
run = para.add_run('\n')
run.font.name = '나눔고딕'
run = para.add_run('공격지표  :  ' + sa_KPI_A + '\n')
run.font.name = '나눔고딕'
run.font.size = Pt(15) # font size

run = para.add_run('수비지표  :  ' + sa_KPI_D + '\n')
run.font.name = '나눔고딕'
run.font.size = Pt(15) # font size

run = para.add_run('성공률지표:  ' + sa_KPI_P + '\n')
run.font.name = '나눔고딕'
run.font.size = Pt(15) # font size

para = doc.add_paragraph()
run = para.add_run('\n\n\n')
run = para.add_run(time)
run.font.name = '나눔고딕'
run.font.size = Pt(20) # font size
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

para = doc.add_paragraph()
run = para.add_run('\n\n\n')
run = para.add_run('용인대학교 산학협력단')
run.font.name = '나눔고딕'
run.bold = True
run.font.size = Pt(20) # font size
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save('평가지표 증명서 자동생성 ' + name + '.docx')

#convert('평가지표 증명서 자동생성 ' + name + '.docx',
#        '평가지표 증명서 자동생성 ' + name + '.pdf')
import re

def convert_to(folder, source, timeout=None):
    args = [libreoffice_exec(), '--headless', '--convert-to', 'pdf', '--outdir', folder, source]

    process = subprocess.run(args, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=timeout)
    filename = re.search('-> (.*?) using filter', process.stdout.decode())

    return filename.group(1)

def libreoffice_exec():
    # TODO: Provide support for more platforms
    if sys.platform == 'darwin':
        return '/usr/bin/soffice'
    return '/usr/bin/soffice'

convert_to('/mnt/d/Work/Python/django/SA_BSP/media/certification/', "평가지표 증명서 자동생성 " + name + '.docx', timeout=15)

#install
#pip install python-docx
#sudo add-apt-repository ppa:libreoffice/pp
#sudo apt-get update
#sudo apt-get upgrade
#sudo apt install libreoffice
#sudo apt install libreoffice-java-common